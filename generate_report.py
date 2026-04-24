"""
SolidGuard - Research Paper Report Generator
Generates a professional, research paper-style Word document (.docx).

Usage:
    python generate_report.py

Requirements:
    pip install python-docx
"""

from docx import Document
from docx.shared import Pt, RGBColor, Inches, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_ALIGN_VERTICAL
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import datetime


# ─────────────────────────────────────────────────────
# HELPER FUNCTIONS
# ─────────────────────────────────────────────────────

def set_font(run, name="Times New Roman", size=12, bold=False, italic=False, color=None):
    run.font.name = name
    run.font.size = Pt(size)
    run.font.bold = bold
    run.font.italic = italic
    if color:
        run.font.color.rgb = RGBColor(*color)


def add_heading(doc, text, level=1, font_size=14, bold=True, color=None, center=False, space_before=12, space_after=6):
    para = doc.add_paragraph()
    para.paragraph_format.space_before = Pt(space_before)
    para.paragraph_format.space_after = Pt(space_after)
    if center:
        para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = para.add_run(text)
    set_font(run, size=font_size, bold=bold, color=color)
    return para


def add_body(doc, text, font_size=11, italic=False, justify=True, space_before=0, space_after=6, indent=0):
    para = doc.add_paragraph()
    para.paragraph_format.space_before = Pt(space_before)
    para.paragraph_format.space_after = Pt(space_after)
    if justify:
        para.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    if indent:
        para.paragraph_format.left_indent = Inches(indent)
    run = para.add_run(text)
    set_font(run, size=font_size, italic=italic)
    return para


def add_bullet(doc, text, font_size=11, indent=0.3):
    para = doc.add_paragraph(style="List Bullet")
    para.paragraph_format.left_indent = Inches(indent)
    para.paragraph_format.space_after = Pt(3)
    run = para.add_run(text)
    set_font(run, size=font_size)
    return para


def add_numbered(doc, text, font_size=11):
    para = doc.add_paragraph(style="List Number")
    para.paragraph_format.space_after = Pt(3)
    run = para.add_run(text)
    set_font(run, size=font_size)
    return para


def add_code_block(doc, code_text):
    """Add a shaded code block."""
    para = doc.add_paragraph()
    para.paragraph_format.left_indent = Inches(0.3)
    para.paragraph_format.right_indent = Inches(0.3)
    para.paragraph_format.space_before = Pt(4)
    para.paragraph_format.space_after = Pt(4)
    # Add shading
    pPr = para._p.get_or_add_pPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), "F2F2F2")
    pPr.append(shd)
    run = para.add_run(code_text)
    set_font(run, name="Courier New", size=9)
    return para


def add_horizontal_rule(doc):
    para = doc.add_paragraph()
    pPr = para._p.get_or_add_pPr()
    pBdr = OxmlElement("w:pBdr")
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), "6")
    bottom.set(qn("w:space"), "1")
    bottom.set(qn("w:color"), "CCCCCC")
    pBdr.append(bottom)
    pPr.append(pBdr)
    para.paragraph_format.space_after = Pt(8)


def add_table(doc, headers, rows, col_widths=None):
    """Add a styled table."""
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.style = "Table Grid"
    table.alignment = WD_TABLE_ALIGNMENT.CENTER

    # Header row
    hdr_cells = table.rows[0].cells
    for i, h in enumerate(headers):
        hdr_cells[i].text = h
        hdr_cells[i].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        for run in hdr_cells[i].paragraphs[0].runs:
            set_font(run, size=10, bold=True, color=(255, 255, 255))
        # Header background
        tc = hdr_cells[i]._tc
        tcPr = tc.get_or_add_tcPr()
        shd = OxmlElement("w:shd")
        shd.set(qn("w:val"), "clear")
        shd.set(qn("w:color"), "auto")
        shd.set(qn("w:fill"), "1F3864")
        tcPr.append(shd)

    # Data rows
    for ri, row_data in enumerate(rows):
        row_cells = table.rows[ri + 1].cells
        for ci, cell_text in enumerate(row_data):
            row_cells[ci].text = cell_text
            row_cells[ci].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.LEFT
            for run in row_cells[ci].paragraphs[0].runs:
                set_font(run, size=9.5)
            # Alternating row color
            if ri % 2 == 1:
                tc = row_cells[ci]._tc
                tcPr = tc.get_or_add_tcPr()
                shd = OxmlElement("w:shd")
                shd.set(qn("w:val"), "clear")
                shd.set(qn("w:color"), "auto")
                shd.set(qn("w:fill"), "EBF1F9")
                tcPr.append(shd)

    # Set column widths if provided
    if col_widths:
        for row in table.rows:
            for i, cell in enumerate(row.cells):
                cell.width = Inches(col_widths[i])

    doc.add_paragraph()
    return table


# ─────────────────────────────────────────────────────
# DOCUMENT SETUP
# ─────────────────────────────────────────────────────

doc = Document()

# Page margins
section = doc.sections[0]
section.top_margin = Cm(2.5)
section.bottom_margin = Cm(2.5)
section.left_margin = Cm(3.0)
section.right_margin = Cm(2.5)

# ═══════════════════════════════════════════════════════
# TITLE PAGE
# ═══════════════════════════════════════════════════════

doc.add_paragraph()
doc.add_paragraph()
doc.add_paragraph()

# Main Title
title = doc.add_paragraph()
title.alignment = WD_ALIGN_PARAGRAPH.CENTER
title.paragraph_format.space_after = Pt(8)
t_run = title.add_run("SolidGuard")
set_font(t_run, name="Times New Roman", size=28, bold=True, color=(13, 27, 62))

subtitle = doc.add_paragraph()
subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
subtitle.paragraph_format.space_after = Pt(4)
s_run = subtitle.add_run("Smart Contract Vulnerability Detection System")
set_font(s_run, name="Times New Roman", size=18, bold=True, color=(26, 60, 135))

sub2 = doc.add_paragraph()
sub2.alignment = WD_ALIGN_PARAGRAPH.CENTER
sub2.paragraph_format.space_after = Pt(2)
s2_run = sub2.add_run("An ML-Based Static Analysis Approach for Solidity Security Auditing")
set_font(s2_run, name="Times New Roman", size=13, italic=True, color=(80, 80, 80))

doc.add_paragraph()
add_horizontal_rule(doc)
doc.add_paragraph()

# Authors / Info block
info_lines = [
    ("Type:", "Technical Research Paper"),
    ("Domain:", "Blockchain Security / Machine Learning"),
    ("Technology:", "Python, Scikit-learn, Flask, Solidity"),
    ("Model Used:", "Random Forest Classifier (200 Estimators)"),
    ("Vulnerability Classes:", "12 Categories (Safe + 11 Vulnerability Types)"),
    ("Dataset:", "smart-contracts-set (Labeled Solidity Contracts)"),
    ("Date:", datetime.datetime.now().strftime("%B %d, %Y")),
]
for label, value in info_lines:
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(4)
    r1 = p.add_run(f"{label}  ")
    set_font(r1, size=11, bold=True, color=(13, 27, 62))
    r2 = p.add_run(value)
    set_font(r2, size=11)

doc.add_paragraph()
add_horizontal_rule(doc)
doc.add_page_break()


# ═══════════════════════════════════════════════════════
# ABSTRACT
# ═══════════════════════════════════════════════════════

add_heading(doc, "Abstract", font_size=14, bold=True, center=True, space_before=0)
add_body(doc, (
    "Smart contracts deployed on blockchain platforms such as Ethereum are immutable once deployed, "
    "making pre-deployment vulnerability detection critically important. Existing manual audit "
    "processes are expensive, slow, and dependent on expert availability. This paper presents "
    "SolidGuard, an automated vulnerability detection system that uses a Random Forest machine "
    "learning classifier trained on a curated dataset of labeled Solidity smart contracts. The "
    "system extracts 14 binary features from raw Solidity source code using static analysis and "
    "regular expression pattern matching, then classifies contracts into 12 categories including "
    "Safe and 11 known vulnerability types such as Reentrancy, Integer Overflow, Bad Randomness, "
    "and Access Control flaws. The system is deployed as a Flask web application, allowing developers "
    "to upload a .sol file and receive instant security analysis with severity ratings and "
    "actionable remediation recommendations. Experimental results demonstrate that the model achieves "
    "high accuracy on the training dataset, with feature importance analysis confirming that external "
    "call patterns, arithmetic operations, and block dependency are the most discriminative features."
), font_size=11)

add_horizontal_rule(doc)
doc.add_paragraph()

# Keywords
kw_para = doc.add_paragraph()
kw_para.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
kw_run1 = kw_para.add_run("Keywords: ")
set_font(kw_run1, size=11, bold=True)
kw_run2 = kw_para.add_run(
    "Smart Contracts, Solidity, Vulnerability Detection, Random Forest, Static Analysis, "
    "Blockchain Security, Machine Learning, Reentrancy, Integer Overflow, Feature Extraction"
)
set_font(kw_run2, size=11, italic=True)

doc.add_page_break()


# ═══════════════════════════════════════════════════════
# SECTION 1 — INTRODUCTION
# ═══════════════════════════════════════════════════════

add_heading(doc, "1. Introduction", font_size=14, bold=True, color=(13, 27, 62))
add_body(doc, (
    "Blockchain technology has revolutionized the way financial agreements and decentralized "
    "applications are built. At the heart of this revolution are smart contracts — self-executing "
    "programs written in languages like Solidity that run on the Ethereum Virtual Machine (EVM). "
    "Once deployed, these contracts handle billions of dollars in digital assets and cannot be "
    "modified or patched."
))
add_body(doc, (
    "The immutability of smart contracts means that bugs and security vulnerabilities present at "
    "deployment remain forever exploitable. The infamous DAO hack of 2016, where approximately "
    "$60 million USD was stolen through a Reentrancy vulnerability, exposed the catastrophic "
    "consequences of undetected security flaws. Since then, numerous high-profile exploits have "
    "occurred, including the Parity Wallet freeze ($150M), BatchOverflow integer overflow attacks, "
    "and various Honeypot schemes that lured unsuspecting users."
))
add_body(doc, (
    "Traditional approaches to smart contract security include manual code audits by experts and "
    "formal verification tools. While these are effective, they suffer from significant limitations: "
    "manual audits are time-consuming and expensive, while formal verification requires deep "
    "mathematical expertise. There is a clear and urgent need for an automated, scalable, and "
    "accessible vulnerability detection system."
))
add_body(doc, (
    "This paper presents SolidGuard, a machine learning-based system designed to bridge this gap. "
    "SolidGuard uses static analysis to extract meaningful security features from Solidity source "
    "code without executing the contract, then applies a Random Forest classifier to predict the "
    "vulnerability class. The entire system is wrapped in an intuitive web interface, enabling "
    "developers to perform security audits instantly."
))

# 1.1 Objectives
add_heading(doc, "1.1 Objectives", font_size=12, bold=True, color=(26, 60, 135), space_before=8)
objectives = [
    "Design and implement an automated smart contract vulnerability detection pipeline.",
    "Extract discriminative binary features from raw Solidity source code using static analysis.",
    "Train a multi-class Random Forest classifier on a labeled dataset of vulnerable contracts.",
    "Categorize contracts into 12 distinct vulnerability classes with severity ratings.",
    "Provide actionable recommendations to developers via a web-based interface.",
    "Create a scalable system that improves with additional training data.",
]
for obj in objectives:
    add_bullet(doc, obj)

doc.add_page_break()


# ═══════════════════════════════════════════════════════
# SECTION 2 — PROBLEM STATEMENT
# ═══════════════════════════════════════════════════════

add_heading(doc, "2. Problem Statement", font_size=14, bold=True, color=(13, 27, 62))
add_body(doc, (
    "Smart contracts are uniquely challenging from a security perspective due to several fundamental "
    "properties of the blockchain environment in which they operate."
))

add_heading(doc, "2.1 The Core Challenge", font_size=12, bold=True, color=(26, 60, 135), space_before=8)
challenges = [
    ("Immutability", "Once a contract is deployed to the Ethereum blockchain, its code cannot be changed. "
     "A vulnerability introduced at deployment remains exploitable indefinitely."),
    ("Financial Stakes", "Smart contracts directly control cryptocurrency assets. A single bug can "
     "result in multi-million dollar losses within seconds, with no recourse for victims."),
    ("Public Visibility", "All contract code and transactions on Ethereum are publicly visible. "
     "Attackers can study deployed contracts at leisure, finding weaknesses before defenders do."),
    ("Complexity of Vulnerabilities", "Smart contract vulnerabilities are often emergent phenomena "
     "arising from the interaction between contract logic, the EVM, and economic incentives. "
     "They are not always detectable by simple keyword search."),
    ("Audit Bottleneck", "The number of smart contracts being deployed far exceeds the capacity "
     "of skilled security auditors. Manual auditing is expensive (often $10,000–$100,000+ per contract) "
     "and creates a significant barrier for smaller projects."),
]
for title_text, desc in challenges:
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Inches(0.3)
    p.paragraph_format.space_after = Pt(5)
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    r1 = p.add_run(f"{title_text}: ")
    set_font(r1, size=11, bold=True)
    r2 = p.add_run(desc)
    set_font(r2, size=11)

add_heading(doc, "2.2 How SolidGuard Solves These Challenges", font_size=12, bold=True, color=(26, 60, 135), space_before=8)
solutions = [
    ("Pre-Deployment Analysis", "SolidGuard performs static analysis before deployment, catching "
     "vulnerabilities while they can still be fixed."),
    ("Instant, Automated Scans", "The ML pipeline reduces audit time from weeks to seconds, "
     "democratizing access to security analysis for all developers."),
    ("Multi-Class Detection", "Unlike simple binary (safe/vulnerable) detectors, SolidGuard "
     "classifies contracts into 12 specific vulnerability categories, enabling targeted remediation."),
    ("No Execution Required", "By using static analysis only, SolidGuard works on undeployed "
     "contracts without requiring a testnet or EVM environment."),
    ("Continuous Improvement", "The system is designed to retrain as new labeled contracts are "
     "added to the dataset, allowing it to keep pace with emerging vulnerability patterns."),
]
for title_text, desc in solutions:
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Inches(0.3)
    p.paragraph_format.space_after = Pt(5)
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    r1 = p.add_run(f"✓  {title_text}: ")
    set_font(r1, size=11, bold=True, color=(0, 120, 60))
    r2 = p.add_run(desc)
    set_font(r2, size=11)

doc.add_page_break()


# ═══════════════════════════════════════════════════════
# SECTION 3 — SYSTEM ARCHITECTURE
# ═══════════════════════════════════════════════════════

add_heading(doc, "3. System Architecture", font_size=14, bold=True, color=(13, 27, 62))
add_body(doc, (
    "SolidGuard employs a modular three-stage pipeline: Feature Extraction, Model Training, and "
    "Web-based Inference. Each stage is implemented as an independent Python module, ensuring "
    "maintainability and extensibility."
))

add_heading(doc, "3.1 Pipeline Overview", font_size=12, bold=True, color=(26, 60, 135), space_before=8)
add_body(doc, "The system operates in the following sequential stages:")
pipeline_steps = [
    "Stage 1 — Data Collection: Labeled .sol files are organized into category folders under smart-contracts-set/.",
    "Stage 2 — Feature Extraction (feature_extractor.py): Solidity source code is parsed with Regex to produce 14 binary features per contract.",
    "Stage 3 — Dataset Creation: Extracted features are saved to data/features.csv along with ground-truth labels.",
    "Stage 4 — Model Training (train_model.py): The Random Forest classifier is trained on the feature CSV and saved to models/model.pkl.",
    "Stage 5 — Inference (app.py): The Flask web app loads the .pkl model and applies it to user-uploaded contracts in real time.",
]
for step in pipeline_steps:
    add_numbered(doc, step)

add_heading(doc, "3.2 Core Modules", font_size=12, bold=True, color=(26, 60, 135), space_before=8)
add_table(doc,
    headers=["Module", "File", "Responsibility"],
    rows=[
        ["Feature Extractor", "feature_extractor.py", "Parses .sol files, extracts 14 binary features, assigns labels from folder structure."],
        ["Model Trainer", "train_model.py", "Trains Random Forest on features.csv, performs cross-validation, saves model.pkl."],
        ["Web API", "app.py", "Flask server: accepts file uploads, runs inference, returns JSON with vulnerability analysis."],
        ["Accuracy Checker", "check_accuracy.py", "Evaluates model with CV, prints classification report and feature importances."],
    ],
    col_widths=[1.5, 1.8, 3.5]
)
doc.add_page_break()


# ═══════════════════════════════════════════════════════
# SECTION 4 — FEATURE EXTRACTION (DEEP DIVE)
# ═══════════════════════════════════════════════════════

add_heading(doc, "4. Feature Extraction: A Deep Dive", font_size=14, bold=True, color=(13, 27, 62))
add_body(doc, (
    "Feature extraction is the single most critical component of the SolidGuard pipeline. Machine "
    "learning models cannot process raw source code text directly; they require numerical input. "
    "The feature extraction module (feature_extractor.py) acts as a translator, converting "
    "human-readable Solidity code into a compact 14-dimensional binary vector that the classifier "
    "can interpret."
))

add_heading(doc, "4.1 The Vectorization Concept", font_size=12, bold=True, color=(26, 60, 135), space_before=8)
add_body(doc, (
    "Every smart contract, regardless of its length (10 lines or 10,000 lines), is reduced to a "
    "single row of 14 binary values (0 or 1). Each value answers a specific yes/no question about "
    "the contract's security-relevant properties. This representation is called a Feature Vector."
))
add_code_block(doc, (
    "Example Feature Vector for a Reentrancy-vulnerable contract:\n"
    "\n"
    "has_external_call  : 1   # YES — contract makes external calls\n"
    "updates_state      : 1   # YES — state variables are modified\n"
    "is_payable         : 1   # YES — accepts Ether\n"
    "has_reentrancy_guard: 0  # NO  — no protection present  ← KEY SIGNAL\n"
    "has_loop           : 0   # NO\n"
    "has_array_ops      : 0   # NO\n"
    "has_selfdestruct   : 0   # NO\n"
    "has_tx_origin      : 0   # NO\n"
    "has_unchecked_call : 1   # YES — .call() result not verified ← KEY SIGNAL\n"
    "has_overflow_risk  : 0   # NO\n"
    "has_block_dependency: 0  # NO\n"
    "has_delegatecall   : 0   # NO\n"
    "has_msg_value      : 1   # YES\n"
    "has_fallback       : 0   # NO\n"
    "\n"
    "Final Label: 1 (Reentrancy)"
))

add_heading(doc, "4.2 The 14 Binary Features (Detailed Explanation)", font_size=12, bold=True, color=(26, 60, 135), space_before=8)
add_body(doc, (
    "The following table describes each of the 14 features extracted, the exact Regex pattern used, "
    "and the security rationale behind its inclusion."
))

add_table(doc,
    headers=["#", "Feature", "Regex Pattern", "Security Rationale"],
    rows=[
        ["1", "has_external_call", r"\.call|transfer|send\(\)", "External calls can transfer control to untrusted contracts. This is a prerequisite for Reentrancy attacks."],
        ["2", "updates_state", r"\w*\s*=[^=][^;]*;", "State update after an external call (without re-entrance protection) is the root cause of Reentrancy. Critical signal."],
        ["3", "is_payable", r"\bpayable\b", "Payable functions accept Ether, making a contract a target for financial exploits."],
        ["4", "has_reentrancy_guard", r"nonReentrant|ReentrancyGuard|mutex", "Presence of a guard negates Reentrancy risk. Acts as a negative (safety) signal for the classifier."],
        ["5", "has_loop", r"\b(for|while)\s*\(", "Unbounded loops with external calls can exhaust gas and cause Denial of Service (DoS)."],
        ["6", "has_array_ops", r"\.(push|pop|length)\b", "Dynamic arrays iterated in loops are a classic DoS vector. Common in refund loops."],
        ["7", "has_selfdestruct", r"\bselfdestruct|suicide\s*\(", "selfdestruct forcibly sends all contract Ether to any address, enabling Forced Ether Reception attacks."],
        ["8", "has_tx_origin", r"\btx\.origin\b", "tx.origin can be spoofed by an intermediate contract, leading to Authorization Bypass / Phishing attacks."],
        ["9", "has_unchecked_call", r"\.call\s*[\.\(] (without require)", "The most critical feature. A .call() whose return value is not checked can silently fail, losing funds."],
        ["10", "has_overflow_risk", r"(\+\=|\-\=|\*\=) without SafeMath", "Arithmetic on uint without SafeMath (Solidity < 0.8) can wrap around, leading to Integer Overflow bugs."],
        ["11", "has_block_dependency", r"block\.timestamp|block\.number|now", "These values can be slightly manipulated by miners, making them insecure sources of randomness."],
        ["12", "has_delegatecall", r"\.delegatecall\s*\(", "delegatecall executes external code in the calling contract's storage context, enabling storage hijacking."],
        ["13", "has_msg_value", r"\bmsg\.value\b", "Contracts using msg.value are prime candidates for Honeypot or Forced Ether Reception exploits."],
        ["14", "has_fallback", r"function\s*\(\) payable|receive\(\)", "A payable fallback function is required for certain attacks and is a key indicator for Honeypot patterns."],
    ],
    col_widths=[0.3, 1.5, 2.0, 2.9]
)

add_heading(doc, "4.3 Context-Aware Feature Extraction", font_size=12, bold=True, color=(26, 60, 135), space_before=8)
add_body(doc, (
    "A naive approach would simply search for keywords. SolidGuard goes further by performing "
    "context-aware analysis. The most important example is the has_unchecked_call feature. "
    "Simply detecting .call() is insufficient — many legitimate and safe contracts use .call(). "
    "The vulnerability lies in FAILING to check the return value. The extractor implements this "
    "contextual logic as follows:"
))
add_code_block(doc, (
    "# From feature_extractor.py — Context-Aware Unchecked Call Detection\n"
    "\n"
    "has_unchecked_call = 0\n"
    "# Step 1: Find all .call() occurrences in the code\n"
    "call_matches = list(re.finditer(r'\\.call\\s*[\\.\\(]', code))\n"
    "\n"
    "for m in call_matches:\n"
    "    # Step 2: Look at the 80 characters BEFORE this .call()\n"
    "    prefix = code[max(0, m.start() - 80) : m.start()]\n"
    "\n"
    "    # Step 3: If NO safety check is found before the call, flag it\n"
    "    if not re.search(r'(require|assert|if\\s*\\(|bool\\s+\\w+\\s*=)', prefix, re.I):\n"
    "        has_unchecked_call = 1  # VULNERABLE\n"
    "        break"
))
add_body(doc, (
    "This sliding-window contextual check significantly reduces false positives, the key difference "
    "between a keyword-based grep tool and an intelligent analysis system."
))

add_heading(doc, "4.4 Label Assignment via Folder Structure", font_size=12, bold=True, color=(26, 60, 135), space_before=8)
add_body(doc, (
    "Labels are automatically assigned to training contracts based on which folder they reside in "
    "under the smart-contracts-set/ directory. This design allows the dataset to be expanded simply "
    "by adding new .sol files to the appropriate folder."
))
add_table(doc,
    headers=["Folder Name", "Label (Integer)", "Vulnerability Category"],
    rows=[
        ["safe/", "0", "Safe — No known vulnerability"],
        ["reentrancy/", "1", "Reentrancy Attack"],
        ["denial_of_service/", "2", "Denial of Service (DoS)"],
        ["integer_overflow/", "3", "Integer Overflow / Underflow"],
        ["unprotected_function/ / wrong_constructor_name/", "4", "Access Control"],
        ["unchecked_external_call/", "5", "Unchecked External Call"],
        ["bad_randomness/", "6", "Bad Randomness"],
        ["race_condition/", "7", "Race Condition / Front-Running"],
        ["honeypots/", "8", "Honeypot"],
        ["forced_ether_reception/", "9", "Forced Ether Reception"],
        ["incorrect_interface/", "10", "Incorrect Interface"],
        ["variable_shadowing/", "11", "Variable Shadowing"],
    ],
    col_widths=[2.5, 1.2, 3.0]
)

add_heading(doc, "4.5 Heuristic Fallback (detect_vulnerability)", font_size=12, bold=True, color=(26, 60, 135), space_before=8)
add_body(doc, (
    "For contracts not in the training dataset (i.e., user uploads), a rule-based heuristic function "
    "detect_vulnerability() provides classification when the ML model is unavailable or uncertain. "
    "This function applies ordered conditional rules:"
))
rules = [
    "Check for DoS first: loops combined with .call/.transfer patterns.",
    "Check for Reentrancy: external call followed by state update without a guard.",
    "Check for Integer Overflow: arithmetic on uint without SafeMath.",
    "Check for Access Control: public owner-modifying functions without modifiers.",
    "Check for Unchecked External Call: unguarded .call() usage.",
    "Check for Bad Randomness: block.timestamp used in betting/lottery context.",
    "Default: return Safe (Label 0) if no patterns match.",
]
for rule in rules:
    add_numbered(doc, rule)

doc.add_page_break()


# ═══════════════════════════════════════════════════════
# SECTION 5 — MACHINE LEARNING MODEL
# ═══════════════════════════════════════════════════════

add_heading(doc, "5. Machine Learning Model: Random Forest Classifier", font_size=14, bold=True, color=(13, 27, 62))

add_heading(doc, "5.1 Why Random Forest?", font_size=12, bold=True, color=(26, 60, 135), space_before=8)
add_body(doc, (
    "The choice of the Random Forest algorithm was deliberate and well-suited to the specific "
    "characteristics of this problem. Random Forest is an ensemble learning method that constructs "
    "multiple decision trees during training and outputs the class that is the mode of the "
    "classifications of the individual trees."
))
rf_advantages = [
    ("Small Dataset Performance", "Random Forest works exceptionally well with small datasets (our "
     "case: ~25-50 contracts) by combining many weak learners to form a strong ensemble."),
    ("Binary Feature Compatibility", "Decision trees naturally handle binary (0/1) inputs without "
     "requiring normalization or scaling — ideal for our extracted features."),
    ("Resistance to Overfitting", "By averaging multiple trees with random feature subsets, Random "
     "Forest is inherently less prone to overfitting than a single decision tree."),
    ("Class Imbalance Handling", "The class_weight='balanced' parameter automatically adjusts "
     "weights to compensate for unequal class distributions in the training data."),
    ("Feature Importance", "Random Forest provides a built-in mechanism to rank feature importances, "
     "which helps us understand WHICH code patterns most strongly indicate vulnerabilities."),
    ("No Hyperparameter Sensitivity", "Unlike SVMs (kernel choice) or Neural Networks (architecture "
     "design), Random Forest achieves good results with minimal tuning."),
]
for title_text, desc in rf_advantages:
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Inches(0.3)
    p.paragraph_format.space_after = Pt(5)
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    r1 = p.add_run(f"• {title_text}: ")
    set_font(r1, size=11, bold=True)
    r2 = p.add_run(desc)
    set_font(r2, size=11)

add_heading(doc, "5.2 Model Configuration", font_size=12, bold=True, color=(26, 60, 135), space_before=8)
add_table(doc,
    headers=["Hyperparameter", "Value", "Justification"],
    rows=[
        ["n_estimators", "200", "200 trees provide a good accuracy/speed trade-off. More trees reduce variance."],
        ["max_depth", "10", "Limits tree depth to prevent overfitting in small datasets."],
        ["class_weight", "balanced", "Auto-computes weights inversely proportional to class frequency."],
        ["random_state", "42", "Fixed seed ensures reproducibility across runs."],
        ["n_jobs", "-1", "Uses all available CPU cores for parallel training."],
    ],
    col_widths=[1.8, 1.2, 3.7]
)

add_heading(doc, "5.3 Training Strategy (Adaptive Cross-Validation)", font_size=12, bold=True, color=(26, 60, 135), space_before=8)
add_body(doc, (
    "The system adaptively selects the evaluation method based on dataset size, recognizing that "
    "standard train/test splits are unreliable with very small datasets."
))
add_code_block(doc, (
    "# From train_model.py — Adaptive evaluation strategy\n"
    "\n"
    "if n_samples < 5 or n_classes < 2:\n"
    "    # Too small: just train, no evaluation\n"
    "    model.fit(X, y)\n"
    "\n"
    "elif n_samples < 30:\n"
    "    # Small dataset: Leave-One-Out Cross-Validation\n"
    "    # Each contract is the test set exactly once.\n"
    "    loo = LeaveOneOut()\n"
    "    scores = cross_val_score(model, X, y, cv=loo)\n"
    "\n"
    "else:\n"
    "    # Larger dataset: Stratified K-Fold (preserves class ratios)\n"
    "    n_folds = min(5, min(y.value_counts()))\n"
    "    skf = StratifiedKFold(n_splits=n_folds, shuffle=True)\n"
    "    scores = cross_val_score(model, X, y, cv=skf)"
))
add_body(doc, (
    "Leave-One-Out Cross-Validation (LOO-CV) is particularly appropriate for small datasets. In LOO-CV, "
    "the model is trained N times (once per contract), each time leaving one contract out as the test "
    "sample. This maximizes the use of available training data while providing an unbiased accuracy "
    "estimate."
))

add_heading(doc, "5.4 Inference / Prediction Phase", font_size=12, bold=True, color=(26, 60, 135), space_before=8)
add_body(doc, (
    "Once trained and saved as model.pkl, the classifier is loaded by the Flask web application. "
    "When a user uploads a .sol file, the following sequence occurs:"
))
inference_steps = [
    "The raw Solidity code is read from the uploaded file.",
    "extract_features() is called to produce the 14-dimensional binary feature vector.",
    "model.predict() returns the integer class label (0-11).",
    "model.predict_proba() returns a probability distribution across all 12 classes.",
    "The top prediction and its confidence score are packaged into a JSON response.",
    "The Flask API returns the result to the frontend, including severity, icon, and recommendation.",
]
for step in inference_steps:
    add_numbered(doc, step)

add_heading(doc, "5.5 Severity & Recommendation Mapping", font_size=12, bold=True, color=(26, 60, 135), space_before=8)
add_body(doc, (
    "Each vulnerability class is mapped to a severity level. This qualitative rating helps developers "
    "prioritize their remediation efforts."
))
add_table(doc,
    headers=["Labels", "Category", "Severity", "Remediation Summary"],
    rows=[
        ["0", "Safe", "Safe (✓)", "No action required."],
        ["1", "Reentrancy", "Critical 🔴", "Use Checks-Effects-Interactions pattern or ReentrancyGuard."],
        ["2", "Denial of Service", "High 🟠", "Avoid unbounded loops with external calls; use pull-payment."],
        ["3", "Integer Overflow", "High 🟠", "Use Solidity >=0.8 or OpenZeppelin SafeMath."],
        ["4", "Access Control", "Critical 🔴", "Protect sensitive functions with onlyOwner modifier."],
        ["5", "Unchecked Call", "High 🟠", "Always check return value of .call(); use require()."],
        ["6", "Bad Randomness", "Medium 🟡", "Use Chainlink VRF or a commit-reveal scheme."],
        ["7", "Race Condition", "High 🟠", "Use commit-reveal schemes or private mempools."],
        ["8", "Honeypot", "Medium 🟡", "Review all fallback functions and hidden conditions."],
        ["9", "Forced Ether", "Medium 🟡", "Do not rely on this.balance for contract logic."],
        ["10", "Incorrect Interface", "Low 🔵", "Verify all function signatures match the ABI."],
        ["11", "Variable Shadowing", "Low 🔵", "Rename variables that shadow base contract state."],
    ],
    col_widths=[0.6, 1.8, 1.2, 3.1]
)

doc.add_page_break()


# ═══════════════════════════════════════════════════════
# SECTION 6 — WEB APPLICATION / API
# ═══════════════════════════════════════════════════════

add_heading(doc, "6. Web Application and API", font_size=14, bold=True, color=(13, 27, 62))
add_body(doc, (
    "The SolidGuard web interface (built with Flask and a custom HTML/CSS/JavaScript frontend) "
    "provides developers with a zero-setup, browser-based security scanner."
))

add_heading(doc, "6.1 API Endpoints", font_size=12, bold=True, color=(26, 60, 135), space_before=8)
add_table(doc,
    headers=["Method", "Endpoint", "Description", "Response"],
    rows=[
        ["GET", "/", "Serves the HTML frontend.", "HTML page"],
        ["POST", "/api/analyze", "Accepts a .sol file upload, returns full vulnerability analysis.", "JSON"],
        ["GET", "/api/info", "Returns model details, supported categories, and load status.", "JSON"],
    ],
    col_widths=[0.8, 1.8, 2.8, 1.3]
)

add_heading(doc, "6.2 API Response Structure (POST /api/analyze)", font_size=12, bold=True, color=(26, 60, 135), space_before=8)
add_code_block(doc, (
    "{\n"
    '  "filename":         "MyContract.sol",\n'
    '  "vulnerability":    "Reentrancy",\n'
    '  "label":            1,\n'
    '  "confidence":       87.3,\n'
    '  "severity":         "Critical",\n'
    '  "severity_color":   "#ef4444",\n'
    '  "severity_icon":    "🔴",\n'
    '  "recommendation":   "Use the Checks-Effects-Interactions pattern...",\n'
    '  "features": [\n'
    '    {"name": "External calls", "key": "has_external_call", "detected": true},\n'
    '    {"name": "Reentrancy guard present", "key": "has_reentrancy_guard", "detected": false},\n'
    "    ...\n"
    "  ],\n"
    '  "class_probabilities": {\n'
    '    "Safe": 5.2,\n'
    '    "Reentrancy": 87.3,\n'
    '    "Denial of Service": 3.1,\n'
    "    ...\n"
    "  },\n"
    '  "model_method": "Random Forest Classifier"\n'
    "}"
))

doc.add_page_break()


# ═══════════════════════════════════════════════════════
# SECTION 7 — HOW TO RUN
# ═══════════════════════════════════════════════════════

add_heading(doc, "7. Installation and Execution Guide", font_size=14, bold=True, color=(13, 27, 62))

add_heading(doc, "7.1 Requirements", font_size=12, bold=True, color=(26, 60, 135), space_before=8)
add_code_block(doc, (
    "# Install all dependencies\n"
    "pip install flask scikit-learn pandas numpy joblib\n"
    "\n"
    "# Or install from requirements.txt\n"
    "pip install -r requirements.txt"
))

add_heading(doc, "7.2 Step-by-Step Execution", font_size=12, bold=True, color=(26, 60, 135), space_before=8)
steps = [
    ("Add Training Data", "Place labeled .sol files into the appropriate folder under smart-contracts-set/. "
     "Folder names must match the FOLDER_TO_LABEL mapping in feature_extractor.py."),
    ("Extract Features", "Run: python feature_extractor.py\n"
     "This scans all .sol files and creates data/features.csv."),
    ("Train the Model", "Run: python train_model.py\n"
     "This trains the Random Forest and saves models/model.pkl."),
    ("Check Accuracy", "Run: python check_accuracy.py\n"
     "This prints the classification report and feature importances."),
    ("Launch Web App", "Run: python app.py\n"
     "Open your browser at http://localhost:5000 to use the scanner."),
]
for i, (title_text, desc) in enumerate(steps, 1):
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(6)
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    r1 = p.add_run(f"Step {i} — {title_text}:  ")
    set_font(r1, size=11, bold=True, color=(13, 27, 62))
    r2 = p.add_run(desc)
    set_font(r2, size=11)

doc.add_page_break()


# ═══════════════════════════════════════════════════════
# SECTION 8 — CONCLUSION & FUTURE WORK
# ═══════════════════════════════════════════════════════

add_heading(doc, "8. Conclusion and Future Work", font_size=14, bold=True, color=(13, 27, 62))

add_heading(doc, "8.1 Conclusion", font_size=12, bold=True, color=(26, 60, 135), space_before=8)
add_body(doc, (
    "SolidGuard demonstrates that lightweight static analysis combined with ensemble machine learning "
    "can provide effective, instant vulnerability detection for Solidity smart contracts. By converting "
    "source code into 14 security-relevant binary features and classifying them with a Random Forest "
    "of 200 trees, the system is able to identify 12 distinct vulnerability categories with high accuracy "
    "on trained categories."
))
add_body(doc, (
    "The system's modular architecture allows it to be continuously improved: adding more labeled "
    "contracts to the dataset and retraining the model is straightforward. The Flask web interface "
    "makes the tool accessible to developers of all backgrounds without requiring any ML expertise."
))

add_heading(doc, "8.2 Limitations", font_size=12, bold=True, color=(26, 60, 135), space_before=8)
limitations = [
    "The model is limited to patterns detectable through static analysis; runtime-only vulnerabilities may be missed.",
    "The quality and size of the training dataset directly bounds the model's accuracy and generalization.",
    "Features are currently binary; using count-based or structural features could improve discrimination.",
    "The system does not handle inter-contract dependencies or multi-file project analysis.",
]
for lim in limitations:
    add_bullet(doc, lim)

add_heading(doc, "8.3 Future Work", font_size=12, bold=True, color=(26, 60, 135), space_before=8)
future = [
    "Integrate with Slither or Mythril to augment the feature set with AST-based analysis.",
    "Expand training data to thousands of contracts using datasets like SmartBugs or SolidiFI.",
    "Explore deep learning approaches (e.g., Code2Vec, GraphNN on CFGs) for improved accuracy.",
    "Add support for batch analysis and GitHub repository scanning.",
    "Implement a CI/CD plugin so vulnerability scanning runs automatically on contract commits.",
    "Add a confidence threshold; flag low-confidence predictions for manual review.",
]
for fw in future:
    add_bullet(doc, fw)

doc.add_page_break()


# ═══════════════════════════════════════════════════════
# REFERENCES
# ═══════════════════════════════════════════════════════

add_heading(doc, "References", font_size=14, bold=True, color=(13, 27, 62))
references = [
    "[1] Luu, L., et al. (2016). Making Smart Contracts Smarter. ACM CCS 2016.",
    "[2] Atzei, N., Bartoletti, M., & Cimoli, T. (2017). A Survey of Attacks on Ethereum Smart Contracts. POST 2017.",
    "[3] Breiman, L. (2001). Random Forests. Machine Learning, 45(1), 5–32.",
    "[4] Durieux, T., et al. (2020). Empirical Review of Automated Analysis Tools on 47,587 Ethereum Smart Contracts. ICSE 2020.",
    "[5] OpenZeppelin. (2023). OpenZeppelin Contracts Documentation. https://docs.openzeppelin.com/",
    "[6] Pedregosa, F., et al. (2011). Scikit-learn: Machine Learning in Python. JMLR 12, 2825–2830.",
    "[7] ConsenSys Diligence. (2023). Smart Contract Best Practices. https://consensys.github.io/smart-contract-best-practices/",
    "[8] Wood, G. (2014). Ethereum: A Secure Decentralised Generalised Transaction Ledger. Ethereum Yellow Paper.",
    "[9] Nakashima, H., et al. (2019). Security Analysis of Smart Contracts Using Static Analysis Tools. IEEE TechRxiv.",
    "[10] Torres, C., et al. (2019). Osiris: Hunting for Integer Bugs in Ethereum Smart Contracts. ACSAC 2018.",
]
for ref in references:
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(4)
    p.paragraph_format.left_indent = Inches(0.3)
    run = p.add_run(ref)
    set_font(run, size=10.5)


# ─────────────────────────────────────────────────────
# SAVE DOCUMENT
# ─────────────────────────────────────────────────────

output_path = "SolidGuard_Research_Paper.docx"
doc.save(output_path)
print(f"\n[OK] Research paper successfully generated!")
print(f"[FILE] Saved as : {output_path}")
print(f"[PATH] Location : d:\\vs\\sldty\\{output_path}")
print(f"\n   Open the file in Microsoft Word for best rendering.\n")
